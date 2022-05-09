from docx import Document
from docx.shared import Pt, Cm

from syllabit import settings
from syllabus.models import Syllabus


def beautify_row(*rows, center: bool):
    for row in rows:
        for cell in row.cells:
            if center:
                cell.paragraphs[0].alignment = 1
            cell.paragraphs[0].runs[0].font.bold = True


def save_docx(syl: Syllabus):
    document = Document()

    # Styles
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Header
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    r = table.cell(0, 0).paragraphs[0].add_run()
    r.add_picture('/home/khandosaly/syllabit/back/media/aitu_logo.png', width=Cm(3.5), height=Cm(2))
    table.cell(0, 1).paragraphs[0].alignment = 2
    table.cell(0, 1).paragraphs[0].add_run('Approved:\n')
    table.cell(0, 1).paragraphs[0].add_run(f'Coordinator: {syl.approvers.filter(type="COORDINATOR").first().user.name}\n')
    table.cell(0, 1).paragraphs[0].add_run(f'Dean: {syl.approvers.filter(type="DEAN").first().user.name}\n')
    table.cell(0, 1).paragraphs[0].add_run('25.04.2022')

    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run('\nSyllabus\n')
    r.bold = True
    r = p.add_run('for the course\n')
    r = p.add_run(f'«{syl.course.name}»\n')
    r.bold = True
    r = p.add_run(f'Academic Year {syl.year}\n')

    # Table
    table = document.add_table(rows=15, cols=2)
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(12)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'

    table.cell(0, 0).text = '1. General information'
    table.cell(0, 0).merge(table.cell(0, 1))
    beautify_row(table.rows[0], center=True)

    table.cell(1, 0).text = 'Major code and title'
    table.cell(1, 1).text = f'{syl.course.name}'

    table.cell(2, 0).text = 'Subject category'
    table.cell(2, 1).text = 'Basic'

    table.cell(3, 0).text = 'Number of Credits'
    table.cell(3, 1).text = f'{syl.course.credits}'

    language_code_dict = {
        'KZ': 'Казахский',
        'RU': 'Русский',
        'EN': 'Английский'
    }
    table.cell(4, 0).text = 'Language of delivery'
    table.cell(4, 1).text = f'{language_code_dict[syl.language]}'

    table.cell(5, 0).text = 'Prerequisites'
    table.cell(5, 1).text = ", ".join([x.name for x in syl.pre_courses.all()])

    table.cell(6, 0).text = 'Postrequisites'
    table.cell(6, 1).text = ", ".join([x.name for x in syl.post_courses.all()])

    table.cell(7, 0).text = 'Instructors'
    table.cell(7, 1).text = ", ".join([x.name for x in syl.teachers.all()])

    table.cell(8, 0).text = '2. Goals, objectives and learning outcomes of the course'
    table.cell(8, 0).merge(table.cell(8, 1))
    beautify_row(table.rows[8], center=True)

    p = table.cell(9, 0).paragraphs[0]
    p.add_run('Course goal(s):\n')
    p = table.cell(9, 0).add_paragraph()
    p.add_run(syl.purpose)
    table.cell(9, 0).merge(table.cell(9, 1))
    beautify_row(table.rows[9], center=False)

    p = table.cell(10, 0).paragraphs[0]
    p.add_run('Course objectives:\n')
    p = table.cell(10, 0).add_paragraph()
    p.add_run(syl.tasks)
    table.cell(10, 0).merge(table.cell(10, 1))
    beautify_row(table.rows[10], center=False)

    p = table.cell(11, 0).paragraphs[0]
    p.add_run('Learning outcomes:\n')
    p = table.cell(11, 0).add_paragraph()
    p.add_run(syl.results)
    table.cell(11, 0).merge(table.cell(11, 1))
    beautify_row(table.rows[11], center=False)

    p = table.cell(12, 0).paragraphs[0]
    p.add_run('3. Course description:\n')
    p = table.cell(12, 0).add_paragraph()
    p.add_run(syl.description)
    table.cell(12, 0).merge(table.cell(12, 1))
    beautify_row(table.rows[12], center=True)

    p = table.cell(13, 0).paragraphs[0]
    p.add_run('4. Course policy: \n')
    p = table.cell(13, 0).add_paragraph()
    p.add_run(syl.policy)
    table.cell(13, 0).merge(table.cell(13, 1))
    beautify_row(table.rows[13], center=True)

    p = table.cell(14, 0).paragraphs[0]
    p.add_run('5. Literature: \n')
    p = table.cell(14, 0).add_paragraph()
    p.add_run(syl.resources)
    table.cell(14, 0).merge(table.cell(14, 1))
    beautify_row(table.rows[14], center=True)

    # Calendar
    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run('\n6. Course calendar')
    r.bold = True
    p = document.add_paragraph()
    r = p.add_run('6.1 Lecture, practical/seminar/laboratory session plans')
    r.bold = True

    table = document.add_table(rows=6, cols=3)
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(4)
    table.columns[2].width = Cm(10)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'

    abbreviations = (
        ('№', 'Abbreviation', 'Meaning'),
        ('1', 'TSIS', 'Teacher-supervised independent work (СРСП)'),
        ('2', 'SIS', 'Students’ independent work (СРС)'),
        ('3', 'IP/GP', 'Individual or group project'),
        ('4', 'LW', 'Laboratory Work'),
        ('5', 'Q', 'Quiz'),
    )

    for i, r in enumerate(abbreviations):
        table.cell(i, 0).text = r[0]
        table.cell(i, 1).text = r[1]
        table.cell(i, 2).text = r[2]

    p = document.add_paragraph()

    # Topics
    table = document.add_table(rows=12, cols=8)
    for i in range(0, 8):
        table.columns[i].width = Cm(1.5)
    table.columns[2].width = Cm(3)
    table.columns[1].width = Cm(7)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'

    topics = [('W №', 'Topic', 'References', 'L', 'P', 'LAB', 'SIS', 'TSIS')]
    for t in syl.topics.order_by('week').all():
        topics.append(
            (str(t.week), t.name, t.materials, str(syl.course.lecture_hours/10), str(syl.course.practice_hours/10),
             '0', str(syl.course.sis_hours/10), str(syl.course.tsis_hours/10))
        )
    topics.append(
        ('', 'Total hours: ',
         str(syl.course.sis_hours+syl.course.tsis_hours+syl.course.lecture_hours+syl.course.practice_hours),
         str(syl.course.lecture_hours), str(syl.course.practice_hours), '0', str(syl.course.sis_hours),
         str(syl.course.tsis_hours))
    )

    for i, r in enumerate(topics):
        for ii, w in enumerate(r):
            table.cell(i, ii).text = r[ii]

    p = document.add_paragraph()
    r = p.add_run('\n7. Student performance evaluation system for the course')
    r.bold = True
    p.alignment = 1
    p = document.add_paragraph()
    r.add_picture(syl.evaluation.photo.path)

    table = document.add_table(rows=12, cols=4)
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(3)
    table.columns[2].width = Cm(3)
    table.columns[3].width = Cm(4)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'

    grades = (
        ('Letter', 'Numerical', 'Percentage', 'Traditional'),
        ('A', '4.0', '95-100', 'Excellent'),
        ('A-', '3.67', '90-94', 'Excellent'),
        ('B+', '3.33', '85-89', 'Good'),
        ('B', '3.0', '80-84', 'Good'),
        ('B-', '2.67', '75-79', 'Good'),
        ('C+', '2.33', '70-74', 'Satisfactory'),
        ('C', '2.0', '65-69', 'Satisfactory'),
        ('C-', '1.67', '60-64', 'Satisfactory'),
        ('D+', '1.33', '55-59', 'Satisfactory'),
        ('D', '1.0', '50-54', 'Satisfactory'),
        ('F', '0', '0-49', 'Fail'),
    )

    for i, r in enumerate(grades):
        for ii, w in enumerate(r):
            table.cell(i, ii).text = r[ii]

    p = document.add_paragraph()
    p.alignment = 1
    r = p.add_run('\n8. Methodological Guidelines')
    r.bold = True
    p = document.add_paragraph()
    r = p.add_run(syl.methods)

    document.save(f'{settings.MEDIA_ROOT}syllabus_{syl.id}.docx')
